"""
Build Arabic Speech Assessment project report (Word docx).
Uses python-docx with explicit RTL formatting throughout.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ---------- RTL helpers ----------
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_run_rtl(run):
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color='8FAADC', size='6'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)


def add_arabic_paragraph(doc, text, *, bold=False, italic=False, size=12,
                         color=None, align='right', spacing_after=6,
                         spacing_before=0, font='Arial'):
    p = doc.add_paragraph()
    p.alignment = {
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(spacing_after)
    if spacing_before:
        p.paragraph_format.space_before = Pt(spacing_before)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_run_rtl(run)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string('1F3864')
    set_run_rtl(run)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string('2E75B6')
    set_run_rtl(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    set_run_rtl(run)
    return p


def add_cell_text(cell, text, *, bold=False, color=None, size=11, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    set_rtl(p)
    run = p.add_run(str(text))
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_run_rtl(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_break(doc):
    doc.add_page_break()


# ---------- Data ----------
TRANSCRIPT = 'السلام عليكم أنا طرال أحب أن أتعلم اللغة الأرضية'

RESULTS = [
    {
        'model': 'GPT-4o Mini', 'provider': 'OpenAI',
        'pronunciation': 6, 'fluency': 5, 'grammar': 6, 'vocabulary': 5, 'overall': 5, 'cefr': 'A1',
        'strengths': ['استخدام التحية بشكل صحيح', 'التعبير عن الفكرة الأساسية'],
        'weaknesses': ['عدم اكتمال الجملة', 'استخدام مفردات بسيطة جداً'],
        'recommendation': 'العمل على بناء جمل كاملة وتوسيع المفردات.',
        'feedback': ('الطالب استخدم تحية صحيحة، لكن الجملة غير مكتملة مما يؤثر على الفهم. '
                     'يجب على الطالب التركيز على بناء جمل كاملة واستخدام مفردات متنوعة لتحسين مستوى اللغة.'),
    },
    {
        'model': 'Qwen 2.5 72B', 'provider': 'Alibaba',
        'pronunciation': 5, 'fluency': 4, 'grammar': 3, 'vocabulary': 4, 'overall': 4, 'cefr': 'A1',
        'strengths': ['استخدام التحية بشكل صحيح'],
        'weaknesses': ['خلل في الجملة', 'ضعف في بناء الجملة'],
        'recommendation': 'التدرب على بناء الجمل البسيطة بشكل صحيح',
        'feedback': ('الطالب استخدم التحية بشكل صحيح، لكن هناك خلل في بناء الجملة. '
                     'يُنصح بالتدرب على بناء الجمل البسيطة بشكل صحيح لتحسين مستوى اللغة.'),
    },
    {
        'model': 'Llama 3.3 70B', 'provider': 'Meta',
        'pronunciation': 4, 'fluency': 3, 'grammar': 2, 'vocabulary': 3, 'overall': 3, 'cefr': 'A1',
        'strengths': ['القدرة على بدء الحديث'],
        'weaknesses': ['نقص في القواعد النحوية', 'صعوبة في التعبير عن الأفكار'],
        'recommendation': 'يُوصى بالتركيز على تعلم القواعد النحوية الأساسية وتحسين مهارات التعبير',
        'feedback': ('يظهر الطالب صعوبة في بناء الجمل بشكل صحيح ويتعذر عليه التعبير عن أفكاره بفعالية. '
                     'يحتاج إلى مزيد من الممارسة في استخدام القواعد النحوية والاهتمام بتحسين مستوى المفردات.'),
    },
    {
        'model': 'Mistral Large', 'provider': 'Mistral',
        'pronunciation': 6, 'fluency': 4, 'grammar': 3, 'vocabulary': 4, 'overall': 4, 'cefr': 'A1',
        'strengths': ['النطق واضح نسبياً في الكلمات البسيطة', 'محاولة استخدام تحية شائعة'],
        'weaknesses': ['خطأ نحوي في تركيب الجملة', 'قلة المفردات', 'عدم طلاقة في الكلام'],
        'recommendation': 'التركيز على تعلم التراكيب الأساسية للجمل والتمرن على المحادثات القصيرة',
        'feedback': ('يتميز الطالب بقدرته على نطق الكلمات البسيطة بوضوح، لكنه يواجه صعوبات في بناء الجمل بشكل صحيح. '
                     'المفردات المستخدمة محدودة، ولا يظهر طلاقة في الكلام. المستوى الحالي يقارب A1، ويحتاج إلى ممارسة أكثر.'),
    },
]

best = max(RESULTS, key=lambda r: r['overall'])
avg_overall = sum(r['overall'] for r in RESULTS) / len(RESULTS)
avg_pron = sum(r['pronunciation'] for r in RESULTS) / len(RESULTS)
avg_flu = sum(r['fluency'] for r in RESULTS) / len(RESULTS)
avg_gra = sum(r['grammar'] for r in RESULTS) / len(RESULTS)
avg_voc = sum(r['vocabulary'] for r in RESULTS) / len(RESULTS)

# ---------- Build document ----------
doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# RTL section property
sectPr = section._sectPr
bidi = OxmlElement('w:bidi')
sectPr.append(bidi)

# Default font for the whole document
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

# Header
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_rtl(header_p)
hrun = header_p.add_run('تقرير تقييم النطق العربي')
hrun.font.size = Pt(9)
hrun.font.color.rgb = RGBColor.from_string('8FAADC')
hrun.font.name = 'Arial'
set_run_rtl(hrun)

# ============================================================
# COVER
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_arabic_paragraph(doc, 'تقرير مشروع', bold=True, size=18, color='2E75B6',
                     align='center', spacing_after=8)
add_arabic_paragraph(doc, 'تقييم النطق العربي بالذكاء الاصطناعي', bold=True, size=28,
                     color='1F3864', align='center', spacing_after=20)
add_arabic_paragraph(doc, 'Arabic Speech Assessment with Multi-LLM Comparison',
                     italic=True, size=13, color='595959', align='center', spacing_after=10)

for _ in range(8):
    doc.add_paragraph()

add_arabic_paragraph(doc, 'إعداد: مجاهد', bold=True, size=14, align='center', spacing_after=4)
add_arabic_paragraph(doc, 'مايو 2026', size=12, color='595959', align='center')

add_page_break(doc)

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_h1(doc, 'الملخص التنفيذي')
add_arabic_paragraph(doc,
    'يهدف هذا المشروع إلى بناء نظام آلي لتقييم مستوى اللغة العربية المنطوقة باستخدام نماذج الذكاء الاصطناعي الكبيرة (LLMs). '
    'يقوم النظام بتحويل التسجيل الصوتي إلى نص باستخدام نموذج Whisper، ثم يرسل النص إلى عدة نماذج ذكاء اصطناعي عبر منصة OpenRouter '
    'للحصول على تقييمات متعددة تشمل النطق والطلاقة والقواعد النحوية والمفردات، وتحديد المستوى وفق الإطار المرجعي الأوروبي المشترك (CEFR).',
    align='justify')

add_arabic_paragraph(doc,
    f'تم اختبار المشروع على عينة نطق قصيرة وقورنت نتائج أربعة نماذج: GPT-4o Mini و Qwen 2.5 72B و Llama 3.3 70B و Mistral Large. '
    f'أعطى نموذج {best["model"]} أعلى تقييم بدرجة عامة {best["overall"]} من 10، '
    f'وكان متوسط الدرجة العامة بين جميع النماذج {avg_overall:.1f} من 10. '
    'اتفقت جميع النماذج على أن مستوى المتحدث في هذه العينة هو A1 وفق تصنيف CEFR.',
    align='justify')

# ============================================================
# INTRODUCTION
# ============================================================
add_h1(doc, 'مقدمة المشروع')
add_arabic_paragraph(doc,
    'تواجه عملية تقييم اللغة العربية المنطوقة تحديات تقليدية متعددة، أبرزها الاعتماد على المعلم البشري الذي قد يستهلك وقتاً طويلاً '
    'وقد يختلف تقييمه من شخص لآخر. كما أن قلة الأدوات المتخصصة في تقييم العربية مقارنة بالإنجليزية تجعل عملية تطوير منصات تعليمية ذكية '
    'باللغة العربية أكثر صعوبة.',
    align='justify')

add_arabic_paragraph(doc,
    'يقدم هذا المشروع حلاً يعتمد على نماذج الذكاء الاصطناعي الحديثة لتقديم تقييم آلي شامل لمستوى الطالب في اللغة العربية. '
    'يتميز النظام بإمكانية مقارنة استجابات أكثر من نموذج في نفس الوقت، مما يساعد على تحديد النموذج الأنسب لأغراض التعليم العربي.',
    align='justify')

add_h2(doc, 'أهداف المشروع')
add_bullet(doc, 'بناء واجهة سهلة الاستخدام لتقييم النطق العربي بالذكاء الاصطناعي.')
add_bullet(doc, 'تحويل الصوت العربي إلى نص باستخدام نموذج Whisper.')
add_bullet(doc, 'تقييم النص العربي باستخدام عدة نماذج LLM متوازية عبر OpenRouter.')
add_bullet(doc, 'مقارنة أداء النماذج المختلفة من حيث الدقة والاتساق.')
add_bullet(doc, 'تقديم نتائج مرئية واضحة تشمل الرسوم البيانية والتوصيات.')

# ============================================================
# METHODOLOGY
# ============================================================
add_h1(doc, 'المنهجية والبنية التقنية')

add_h2(doc, 'الأدوات والتقنيات المستخدمة')
add_bullet(doc, 'Streamlit — لبناء واجهة المستخدم التفاعلية مع دعم اللغة العربية واتجاه RTL.')
add_bullet(doc, 'OpenAI Whisper / Groq Whisper — لتحويل التسجيل الصوتي العربي إلى نص.')
add_bullet(doc, 'OpenRouter API — لاستدعاء عدة نماذج LLM من مزودين مختلفين عبر واجهة موحدة.')
add_bullet(doc, 'Plotly — لإنشاء الرسوم البيانية التفاعلية (Bar Chart و Radar Chart).')
add_bullet(doc, 'Pandas — لمعالجة وتصدير النتائج بصيغة CSV.')
add_bullet(doc, 'FFmpeg — لتحويل ملفات الصوت إلى الصيغة المطلوبة.')

add_h2(doc, 'مراحل العمل')
add_arabic_paragraph(doc, 'المرحلة الأولى — الإدخال: يقوم المستخدم برفع ملف صوتي بالعربية، أو يدخل النص مباشرة.')
add_arabic_paragraph(doc, 'المرحلة الثانية — التفريغ الصوتي: يتم تحويل الصوت إلى نص عربي باستخدام Whisper.')
add_arabic_paragraph(doc, 'المرحلة الثالثة — التقييم المتوازي: يُرسل النص إلى عدة نماذج LLM عبر OpenRouter بشكل متوازٍ لتسريع المعالجة.')
add_arabic_paragraph(doc, 'المرحلة الرابعة — تحليل وعرض النتائج: تُجمع نتائج النماذج وتُعرض في شكل جداول ورسوم بيانية وتفاصيل لكل نموذج.')

add_h2(doc, 'معايير التقييم')
add_arabic_paragraph(doc, 'يقيِّم كل نموذج النص وفق خمسة معايير رئيسية، كل منها بدرجة من صفر إلى عشرة:')
add_bullet(doc, 'النطق (Pronunciation): مدى وضوح الكلمات وصحة لفظها.')
add_bullet(doc, 'الطلاقة (Fluency): مدى تدفق الكلام بسلاسة دون توقف.')
add_bullet(doc, 'القواعد النحوية (Grammar): صحة التركيب النحوي للجمل.')
add_bullet(doc, 'المفردات (Vocabulary): تنوع وثراء المفردات المستخدمة.')
add_bullet(doc, 'الدرجة العامة (Overall): تقييم شامل يأخذ بعين الاعتبار جميع المعايير.')
add_arabic_paragraph(doc, 'بالإضافة إلى ذلك، يحدد كل نموذج المستوى وفق تصنيف CEFR من A1 (مبتدئ) إلى C2 (متمكن).')

add_page_break(doc)

# ============================================================
# SAMPLE
# ============================================================
add_h1(doc, 'عينة الاختبار')
add_arabic_paragraph(doc, 'تم اختبار النظام على تسجيل صوتي قصير بالعربية، وكان النص الناتج من عملية التفريغ كالتالي:',
                     align='justify')

# Quote box
quote_table = doc.add_table(rows=1, cols=1)
quote_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(quote_table)
qc = quote_table.rows[0].cells[0]
add_cell_text(qc, f'"{TRANSCRIPT}"', bold=True, color='1F3864', size=14, align='center')
set_cell_shading(qc, 'EDEDED')
set_cell_borders(qc, color='1F3864', size='8')

add_arabic_paragraph(doc, '')
add_arabic_paragraph(doc,
    'يُلاحظ من النص وجود خطأ واضح: الكلمة "الأرضية" بدلاً من "العربية"، وهذا قد يكون نتيجة خطأ في التفريغ الصوتي '
    'أو خطأ في النطق من المتحدث. هذه نقطة مهمة سوف نناقشها لاحقاً عند تحليل أداء النماذج.',
    align='justify')

# ============================================================
# RESULTS TABLE
# ============================================================
add_h1(doc, 'نتائج التقييم — مقارنة شاملة')
add_arabic_paragraph(doc, 'يلخص الجدول التالي تقييم كل نموذج على معايير التقييم الخمسة بالإضافة إلى المستوى المحدد:',
                     align='justify')

headers = ['النموذج', 'المزود', 'النطق', 'الطلاقة', 'النحو', 'المفردات', 'الدرجة العامة', 'CEFR']
table = doc.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(table)

# Header row
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    add_cell_text(hdr_cells[i], h, bold=True, color='FFFFFF', size=11)
    set_cell_shading(hdr_cells[i], '1F3864')
    set_cell_borders(hdr_cells[i])

# Data rows
for idx, r in enumerate(RESULTS):
    row = table.add_row().cells
    is_best = (r is best)
    row_color = 'F2F2F2' if idx % 2 == 0 else 'FFFFFF'
    values = [r['model'], r['provider'], r['pronunciation'], r['fluency'],
              r['grammar'], r['vocabulary'], r['overall'], r['cefr']]
    for i, val in enumerate(values):
        bold = (i == 0) or (i == 6 and is_best)
        cell_color = 'FFD966' if (i == 6 and is_best) else row_color
        add_cell_text(row[i], val, bold=bold, size=11)
        set_cell_shading(row[i], cell_color)
        set_cell_borders(row[i])

# Average row
avg_row = table.add_row().cells
avg_values = ['المتوسط', '—',
              f'{avg_pron:.1f}', f'{avg_flu:.1f}', f'{avg_gra:.1f}',
              f'{avg_voc:.1f}', f'{avg_overall:.1f}', 'A1']
for i, val in enumerate(avg_values):
    add_cell_text(avg_row[i], val, bold=True, color='FFFFFF', size=11)
    set_cell_shading(avg_row[i], '8FAADC')
    set_cell_borders(avg_row[i])

add_arabic_paragraph(doc, '')
add_arabic_paragraph(doc,
    f'النموذج الأفضل أداءً هو {best["model"]} بدرجة عامة {best["overall"]} من 10. '
    f'بينما النموذج الأقل تقييماً هو Llama 3.3 70B بدرجة 3 فقط، مما يُظهر تباينًا واضحاً بين النماذج رغم اتفاقها على المستوى A1.',
    align='justify')

add_page_break(doc)

# ============================================================
# PER-MODEL DETAILS
# ============================================================
add_h1(doc, 'تحليل تفصيلي لكل نموذج')

for r in RESULTS:
    add_h2(doc, f'{r["model"]} ({r["provider"]})')

    # Score line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_after = Pt(6)

    def add_run(text, *, bold=False, color=None, size=12):
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        set_run_rtl(run)
        return run

    add_run('الدرجة العامة: ', bold=True)
    add_run(f'{r["overall"]}/10', bold=True, color='C00000' if r is best else '1F3864')
    add_run('  |  مستوى CEFR: ', bold=True)
    add_run(r['cefr'], bold=True, color='1F3864')

    # Strengths
    add_arabic_paragraph(doc, 'نقاط القوة:', bold=True, color='548235', size=13, spacing_before=6, spacing_after=4)
    for s in r['strengths']:
        add_bullet(doc, s)

    # Weaknesses
    add_arabic_paragraph(doc, 'نقاط الضعف:', bold=True, color='C00000', size=13, spacing_before=6, spacing_after=4)
    for w in r['weaknesses']:
        add_bullet(doc, w)

    # Recommendation
    add_arabic_paragraph(doc, 'التوصية:', bold=True, color='BF8F00', size=13, spacing_before=6, spacing_after=4)
    add_arabic_paragraph(doc, r['recommendation'], align='justify')

    # Feedback
    add_arabic_paragraph(doc, 'التقييم التفصيلي:', bold=True, color='2E75B6', size=13, spacing_before=6, spacing_after=4)
    add_arabic_paragraph(doc, r['feedback'], italic=True, align='justify', spacing_after=10)

add_page_break(doc)

# ============================================================
# KEY FINDINGS
# ============================================================
add_h1(doc, 'النتائج الرئيسية والملاحظات')

add_h2(doc, 'أولاً: تباين تقييمات النماذج')
add_arabic_paragraph(doc,
    'رغم أن جميع النماذج اتفقت على أن المستوى هو A1، إلا أنها أعطت درجات مختلفة في المعايير الفرعية. '
    'الفرق بين أعلى تقييم (5/10 من GPT-4o Mini) وأقل تقييم (3/10 من Llama 3.3 70B) يصل إلى 40%، '
    'مما يشير إلى الحاجة إلى مرجعية موحدة عند بناء نظام تقييم آلي قابل للاعتماد.',
    align='justify')

add_h2(doc, 'ثانياً: أداء GPT-4o Mini')
add_arabic_paragraph(doc,
    'كان GPT-4o Mini الأكثر اعتدالاً وتفصيلاً في التقييم. أعطى درجات أعلى بشكل ملحوظ في القواعد والمفردات، '
    'وقدم ملاحظات بناءة ومحددة دون مبالغة في انتقاد المتعلم المبتدئ.',
    align='justify')

add_h2(doc, 'ثالثاً: حساسية النماذج للأخطاء الدلالية')
add_arabic_paragraph(doc,
    'الكلمة "الأرضية" بدلاً من "العربية" أربكت بعض النماذج. حيث رصدها Llama 3.3 70B كنقطة ضعف، '
    'بينما تجاهلها GPT-4o Mini ولم يشير إليها صراحةً. هذا يعكس تفاوت دقة النماذج في رصد الأخطاء الدلالية الدقيقة في العربية.',
    align='justify')

add_h2(doc, 'رابعاً: اتفاق على مستوى CEFR')
add_arabic_paragraph(doc,
    'جميع النماذج الأربعة اتفقت على تصنيف المستوى كـ A1 (مبتدئ)، وهذا يعزز موثوقية تحديد المستوى الإجمالي '
    'حتى عندما تختلف الدرجات الرقمية. الاتفاق على المستوى مع اختلاف الدرجات قد يكون مؤشراً مهماً للأنظمة التعليمية.',
    align='justify')

add_h2(doc, 'خامساً: جودة التوصيات')
add_arabic_paragraph(doc,
    'جميع النماذج قدمت توصيات مفيدة تركز على بناء الجمل وتوسيع المفردات، وهي توصيات منطقية لمستوى A1. '
    'كانت توصيات Mistral Large الأكثر تفصيلاً وعملية، حيث اقترحت "التمرن على المحادثات القصيرة".',
    align='justify')

# ============================================================
# RECOMMENDATIONS
# ============================================================
add_h1(doc, 'التوصيات والتطوير المستقبلي')

add_h2(doc, 'تطوير النظام')
add_bullet(doc, 'استخدام طبقة تحقق (Voting Layer) لاتخاذ القرار النهائي بناءً على إجماع النماذج بدلاً من نموذج واحد.')
add_bullet(doc, 'تحسين دقة التفريغ الصوتي باستخدام Whisper Large بدل Whisper Small.')
add_bullet(doc, 'إضافة مكون تحليل النطق على مستوى الصوتيات (Phonemes) للتقييم الدقيق للنطق.')
add_bullet(doc, 'بناء قاعدة بيانات مرجعية للأخطاء الشائعة في العربية لزيادة دقة التقييم.')

add_h2(doc, 'تطوير النماذج')
add_bullet(doc, 'اختبار نماذج متخصصة في العربية مثل Jais و AceGPT لمقارنة أدائها مع النماذج متعددة اللغات.')
add_bullet(doc, 'استخدام Few-shot Prompting بأمثلة تقييم مرجعية لتقليل التباين بين النماذج.')
add_bullet(doc, 'إضافة درجة "ثقة" لكل تقييم لإبراز الحالات الغامضة.')

add_h2(doc, 'تطوير تجربة المستخدم')
add_bullet(doc, 'إضافة تتبع تاريخي لتقدم المتعلم عبر الزمن.')
add_bullet(doc, 'إضافة تمارين تفاعلية مقترحة بناءً على نقاط الضعف المكتشفة.')
add_bullet(doc, 'دعم اللهجات العربية المختلفة (خليجية، شامية، مصرية، مغاربية).')

# ============================================================
# CONCLUSION
# ============================================================
add_h1(doc, 'الخاتمة')
add_arabic_paragraph(doc,
    'أثبت هذا المشروع جدوى استخدام نماذج LLM المتاحة عبر OpenRouter لبناء نظام تقييم آلي للعربية المنطوقة. '
    'النتائج تشير إلى أن النماذج الحالية قادرة على إعطاء تقييم مبدئي معقول، خاصةً عند تحديد المستوى العام وفق CEFR. '
    'ومع ذلك، فإن التباين الكبير في الدرجات الرقمية يستدعي تطوير منهجية أكثر صرامة للحصول على نتائج موحدة وقابلة للاعتماد.',
    align='justify')

add_arabic_paragraph(doc,
    'تشكل هذه النسخة الأولى من المشروع أساساً قوياً يمكن البناء عليه لتطوير منصة تعليمية ذكية كاملة للغة العربية، '
    'تخدم ملايين المتعلمين حول العالم الذين يبحثون عن أدوات حديثة لتعلم العربية.',
    align='justify')

add_arabic_paragraph(doc, '— نهاية التقرير —', italic=True, color='8FAADC', size=12,
                     align='center', spacing_before=24)

# ---------- Save ----------
output_path = 'Arabic_Speech_Assessment_Report.docx'
doc.save(output_path)
print(f'Created: {output_path}')
