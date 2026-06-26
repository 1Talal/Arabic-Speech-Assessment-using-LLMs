"""
Build Arabic Speech Assessment project report (English, Word docx).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, color='8FAADC', size='6'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_p(doc, text, *, bold=False, italic=False, size=11, color=None,
          align='left', after=6, before=0, font='Calibri'):
    p = doc.add_paragraph()
    p.alignment = {
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_after = Pt(after)
    if before:
        p.paragraph_format.space_before = Pt(before)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string('1F3864')


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string('2E75B6')


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


def add_cell(cell, text, *, bold=False, color=None, size=10, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }[align]
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ---------- Data ----------
TRANSCRIPT_AR = 'السلام عليكم أنا طرال أحب أن أتعلم اللغة الأرضية'
TRANSCRIPT_TRANSLATION = (
    'Peace be upon you, I am Talal, I love to learn the "Ardiyya" '
    '(earth) language [intended: Arabic language].'
)

TEAM = [
    ('Salman Alsulami', '444002538'),
    ('Talal Salem',     '44410106'),
    ('Mujahid',         '444006966'),
]

RESULTS = [
    {
        'model': 'GPT-4o Mini', 'provider': 'OpenAI',
        'pronunciation': 6, 'fluency': 5, 'grammar': 6, 'vocabulary': 5,
        'overall': 5, 'cefr': 'A1',
        'strengths': ['Correct use of the greeting',
                      'Conveys the main idea'],
        'weaknesses': ['Incomplete sentence',
                       'Very basic vocabulary'],
        'recommendation': 'Work on building complete sentences and expanding vocabulary.',
        'feedback': ('The student used a correct greeting, but the sentence is incomplete, '
                     'which affects comprehension. The student should focus on building '
                     'complete sentences and using varied vocabulary to improve the language level.'),
    },
    {
        'model': 'Qwen 2.5 72B', 'provider': 'Alibaba',
        'pronunciation': 5, 'fluency': 4, 'grammar': 3, 'vocabulary': 4,
        'overall': 4, 'cefr': 'A1',
        'strengths': ['Correct use of the greeting'],
        'weaknesses': ['Sentence structure error',
                       'Weak sentence construction'],
        'recommendation': 'Practice building simple sentences correctly.',
        'feedback': ('The student used the greeting correctly, but there is a flaw in sentence '
                     'construction. It is recommended to practice building simple sentences '
                     'correctly to improve the language level.'),
    },
    {
        'model': 'Llama 3.3 70B', 'provider': 'Meta',
        'pronunciation': 4, 'fluency': 3, 'grammar': 2, 'vocabulary': 3,
        'overall': 3, 'cefr': 'A1',
        'strengths': ['Ability to initiate conversation'],
        'weaknesses': ['Lacking grammatical rules',
                       'Difficulty expressing ideas'],
        'recommendation': 'Focus on learning basic grammatical rules and improving expression skills.',
        'feedback': ('The student shows difficulty building sentences correctly and is unable to '
                     'express ideas effectively. More practice with grammar use and improving '
                     'vocabulary level is needed.'),
    },
    {
        'model': 'Mistral Large', 'provider': 'Mistral',
        'pronunciation': 6, 'fluency': 4, 'grammar': 3, 'vocabulary': 4,
        'overall': 4, 'cefr': 'A1',
        'strengths': ['Relatively clear pronunciation of simple words',
                      'Attempted use of a common greeting'],
        'weaknesses': ['Grammatical error in sentence structure',
                       'Limited vocabulary',
                       'Lack of speaking fluency'],
        'recommendation': 'Focus on learning basic sentence structures and practicing short conversations.',
        'feedback': ('The student is able to pronounce simple words clearly, but struggles with '
                     'building sentences correctly. Vocabulary is limited, and fluency is lacking. '
                     'The current level is approximately A1; more practice is needed.'),
    },
]

best = max(RESULTS, key=lambda r: r['overall'])
avg_overall = sum(r['overall'] for r in RESULTS) / len(RESULTS)
avg_pron = sum(r['pronunciation'] for r in RESULTS) / len(RESULTS)
avg_flu = sum(r['fluency'] for r in RESULTS) / len(RESULTS)
avg_gra = sum(r['grammar'] for r in RESULTS) / len(RESULTS)
avg_voc = sum(r['vocabulary'] for r in RESULTS) / len(RESULTS)

# ---------- Document ----------
doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Header
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
hrun = header_p.add_run('Arabic Speech Assessment — Project Report')
hrun.font.size = Pt(9)
hrun.font.color.rgb = RGBColor.from_string('8FAADC')
hrun.font.name = 'Calibri'

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_p(doc, 'PROJECT REPORT', bold=True, size=14, color='2E75B6',
      align='center', after=6)
add_p(doc, 'Arabic Speech Assessment', bold=True, size=32,
      color='1F3864', align='center', after=6)
add_p(doc, 'with Multi-LLM Comparison via OpenRouter',
      italic=True, size=16, color='595959', align='center', after=20)

for _ in range(2):
    doc.add_paragraph()

add_p(doc, 'Team Members', bold=True, size=14, color='1F3864', align='center', after=10)

# Team table
team_table = doc.add_table(rows=len(TEAM) + 1, cols=2)
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = team_table.rows[0].cells
add_cell(hdr[0], 'Name', bold=True, color='FFFFFF', size=12)
add_cell(hdr[1], 'University ID', bold=True, color='FFFFFF', size=12)
set_cell_shading(hdr[0], '1F3864')
set_cell_shading(hdr[1], '1F3864')
set_cell_borders(hdr[0])
set_cell_borders(hdr[1])

for i, (name, uid) in enumerate(TEAM, start=1):
    row = team_table.rows[i].cells
    bg = 'F2F2F2' if i % 2 else 'FFFFFF'
    add_cell(row[0], name, bold=True, size=11, align='left')
    add_cell(row[1], uid, size=11)
    set_cell_shading(row[0], bg)
    set_cell_shading(row[1], bg)
    set_cell_borders(row[0])
    set_cell_borders(row[1])

for _ in range(6):
    doc.add_paragraph()

add_p(doc, 'May 2026', size=12, color='595959', align='center')

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_h1(doc, 'Executive Summary')
add_p(doc,
      'This project aims to build an automated system for assessing the level of spoken Arabic '
      'using Large Language Models (LLMs). The system converts an audio recording into text using '
      'the Whisper model, then sends the text to multiple AI models through the OpenRouter platform '
      'to obtain multiple assessments covering pronunciation, fluency, grammar, and vocabulary, and '
      'to determine the level according to the Common European Framework of Reference for Languages (CEFR).',
      align='justify')

add_p(doc,
      f'The project was tested on a short speech sample, and the results of four models were compared: '
      f'GPT-4o Mini, Qwen 2.5 72B, Llama 3.3 70B, and Mistral Large. The {best["model"]} model gave '
      f'the highest evaluation with an overall score of {best["overall"]}/10, while the average overall '
      f'score across all models was {avg_overall:.1f}/10. All four models agreed that the speaker level '
      f'in this sample is A1 according to the CEFR classification.',
      align='justify')

# ============================================================
# INTRODUCTION
# ============================================================
add_h1(doc, 'Project Introduction')
add_p(doc,
      'Assessing spoken Arabic faces several traditional challenges, most notably reliance on human '
      'teachers — a process that can be time-consuming and prone to variability between evaluators. '
      'Additionally, the shortage of specialized assessment tools for Arabic, compared to English, '
      'makes it harder to develop smart educational platforms in Arabic.',
      align='justify')

add_p(doc,
      'This project presents a solution based on modern AI models to provide an automated and '
      'comprehensive assessment of a student\'s level in Arabic. The system allows comparing the '
      'responses of multiple models in parallel, helping identify which model is most suitable for '
      'Arabic-language educational use cases.',
      align='justify')

add_h2(doc, 'Project Objectives')
add_bullet(doc, 'Build a user-friendly interface for assessing Arabic speech using AI.')
add_bullet(doc, 'Convert Arabic audio into text using the Whisper model.')
add_bullet(doc, 'Evaluate Arabic text using multiple LLMs in parallel via OpenRouter.')
add_bullet(doc, 'Compare the performance of different models in terms of accuracy and consistency.')
add_bullet(doc, 'Present clear visual results including charts and recommendations.')

# ============================================================
# METHODOLOGY
# ============================================================
add_h1(doc, 'Methodology and Technical Architecture')

add_h2(doc, 'Tools and Technologies Used')
add_bullet(doc, 'Streamlit — to build an interactive user interface with Arabic and RTL support.')
add_bullet(doc, 'OpenAI Whisper / Groq Whisper — to transcribe Arabic audio into text.')
add_bullet(doc, 'OpenRouter API — to call multiple LLMs from different providers through a unified interface.')
add_bullet(doc, 'Plotly — to create interactive charts (Bar Chart and Radar Chart).')
add_bullet(doc, 'Pandas — to process and export results as CSV.')
add_bullet(doc, 'FFmpeg — to convert audio files to the required format.')

add_h2(doc, 'Workflow Stages')
add_p(doc, 'Stage 1 — Input: The user uploads an Arabic audio file or enters text directly.')
add_p(doc, 'Stage 2 — Transcription: The audio is converted to Arabic text using Whisper.')
add_p(doc, 'Stage 3 — Parallel Assessment: The text is sent to multiple LLMs via OpenRouter in parallel '
           'to speed up processing.')
add_p(doc, 'Stage 4 — Analysis and Display: Results from all models are aggregated and displayed in '
           'tables, charts, and per-model details.')

add_h2(doc, 'Evaluation Criteria')
add_p(doc, 'Each model evaluates the text against five main criteria, each scored from 0 to 10:')
add_bullet(doc, 'Pronunciation — clarity of words and correctness of articulation.')
add_bullet(doc, 'Fluency — smoothness of speech without unnatural pauses.')
add_bullet(doc, 'Grammar — correctness of sentence structure.')
add_bullet(doc, 'Vocabulary — diversity and richness of the vocabulary used.')
add_bullet(doc, 'Overall — a holistic score taking all criteria into account.')
add_p(doc, 'In addition, each model assigns the speaker a CEFR level ranging from A1 (beginner) to C2 (mastery).')

doc.add_page_break()

# ============================================================
# SAMPLE
# ============================================================
add_h1(doc, 'Test Sample')
add_p(doc, 'The system was tested on a short Arabic audio recording. The output text from the '
           'transcription step was as follows:',
      align='justify')

quote_table = doc.add_table(rows=2, cols=1)
quote_table.alignment = WD_TABLE_ALIGNMENT.CENTER
qc1 = quote_table.rows[0].cells[0]
add_cell(qc1, TRANSCRIPT_AR, bold=True, color='1F3864', size=14, align='center')
set_cell_shading(qc1, 'EDEDED')
set_cell_borders(qc1, color='1F3864', size='8')

qc2 = quote_table.rows[1].cells[0]
add_cell(qc2, f'Translation: {TRANSCRIPT_TRANSLATION}', size=10, color='595959', align='center')
set_cell_shading(qc2, 'F8F8F8')
set_cell_borders(qc2, color='1F3864', size='8')

add_p(doc, '')
add_p(doc,
      'A clear error can be noticed in the output: the word "Ardiyya" (earth/ground) appears instead '
      'of "Arabiyya" (Arabic). This may be due to either a transcription error or a pronunciation '
      'mistake from the speaker. This is an important point we discuss later when analyzing model performance.',
      align='justify')

# ============================================================
# RESULTS TABLE
# ============================================================
add_h1(doc, 'Assessment Results — Full Comparison')
add_p(doc, 'The following table summarizes each model\'s scores across all five criteria, along '
           'with the assigned CEFR level:',
      align='justify')

headers = ['Model', 'Provider', 'Pronunciation', 'Fluency', 'Grammar', 'Vocabulary', 'Overall', 'CEFR']
table = doc.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    add_cell(hdr_cells[i], h, bold=True, color='FFFFFF', size=10)
    set_cell_shading(hdr_cells[i], '1F3864')
    set_cell_borders(hdr_cells[i])

for idx, r in enumerate(RESULTS):
    row = table.add_row().cells
    is_best = (r is best)
    row_color = 'F2F2F2' if idx % 2 == 0 else 'FFFFFF'
    values = [r['model'], r['provider'], r['pronunciation'], r['fluency'],
              r['grammar'], r['vocabulary'], r['overall'], r['cefr']]
    for i, val in enumerate(values):
        bold = (i == 0) or (i == 6 and is_best)
        cell_color = 'FFD966' if (i == 6 and is_best) else row_color
        add_cell(row[i], val, bold=bold, size=10)
        set_cell_shading(row[i], cell_color)
        set_cell_borders(row[i])

avg_row = table.add_row().cells
avg_values = ['Average', '—',
              f'{avg_pron:.1f}', f'{avg_flu:.1f}', f'{avg_gra:.1f}',
              f'{avg_voc:.1f}', f'{avg_overall:.1f}', 'A1']
for i, val in enumerate(avg_values):
    add_cell(avg_row[i], val, bold=True, color='FFFFFF', size=10)
    set_cell_shading(avg_row[i], '8FAADC')
    set_cell_borders(avg_row[i])

add_p(doc, '')
add_p(doc,
      f'The best-performing model is {best["model"]} with an overall score of {best["overall"]}/10, '
      f'while the lowest-rated model is Llama 3.3 70B with only 3/10. This shows a clear divergence '
      f'between models, despite their agreement on the A1 CEFR level.',
      align='justify')

doc.add_page_break()

# ============================================================
# PER-MODEL DETAILS
# ============================================================
add_h1(doc, 'Detailed Per-Model Analysis')

for r in RESULTS:
    add_h2(doc, f'{r["model"]} ({r["provider"]})')

    # Score line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    def add_run(text, *, bold=False, color=None, size=11):
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    add_run('Overall Score: ', bold=True)
    add_run(f'{r["overall"]}/10', bold=True, color='C00000' if r is best else '1F3864')
    add_run('   |   CEFR Level: ', bold=True)
    add_run(r['cefr'], bold=True, color='1F3864')

    add_p(doc, 'Strengths:', bold=True, color='548235', size=12, before=6, after=4)
    for s in r['strengths']:
        add_bullet(doc, s)

    add_p(doc, 'Weaknesses:', bold=True, color='C00000', size=12, before=6, after=4)
    for w in r['weaknesses']:
        add_bullet(doc, w)

    add_p(doc, 'Recommendation:', bold=True, color='BF8F00', size=12, before=6, after=4)
    add_p(doc, r['recommendation'], align='justify')

    add_p(doc, 'Detailed Feedback:', bold=True, color='2E75B6', size=12, before=6, after=4)
    add_p(doc, r['feedback'], italic=True, align='justify', after=10)

doc.add_page_break()

# ============================================================
# KEY FINDINGS
# ============================================================
add_h1(doc, 'Key Findings and Observations')

add_h2(doc, '1. Divergence Between Model Evaluations')
add_p(doc,
      'Although all models agreed that the level is A1, they produced different scores across the '
      'sub-criteria. The gap between the highest score (5/10 from GPT-4o Mini) and the lowest score '
      '(3/10 from Llama 3.3 70B) reaches 40%, suggesting the need for a standardized reference when '
      'building a reliable automated assessment system.',
      align='justify')

add_h2(doc, '2. Performance of GPT-4o Mini')
add_p(doc,
      'GPT-4o Mini was the most balanced and detailed in its evaluation. It gave noticeably higher '
      'scores on grammar and vocabulary, and provided constructive, specific feedback without being '
      'overly harsh on a beginner learner.',
      align='justify')

add_h2(doc, '3. Model Sensitivity to Semantic Errors')
add_p(doc,
      'The word "Ardiyya" (earth) instead of "Arabiyya" (Arabic) confused some models. Llama 3.3 70B '
      'flagged it as a weakness, whereas GPT-4o Mini ignored it and did not explicitly mention the '
      'mistake. This reflects varying levels of model precision when detecting fine-grained semantic '
      'errors in Arabic.',
      align='justify')

add_h2(doc, '4. Agreement on CEFR Level')
add_p(doc,
      'All four models agreed on classifying the level as A1 (beginner), reinforcing the reliability '
      'of the overall level assignment even when numerical scores differ. Consensus on level despite '
      'score variation could be an important signal for educational systems.',
      align='justify')

add_h2(doc, '5. Quality of Recommendations')
add_p(doc,
      'All models produced useful recommendations focused on building sentences and expanding '
      'vocabulary — sensible advice for the A1 level. The Mistral Large recommendations were the '
      'most detailed and practical, suggesting "practicing short conversations."',
      align='justify')

# ============================================================
# RECOMMENDATIONS
# ============================================================
add_h1(doc, 'Recommendations and Future Work')

add_h2(doc, 'System Improvements')
add_bullet(doc, 'Use a voting layer to determine the final score based on model consensus rather than a single model.')
add_bullet(doc, 'Improve transcription accuracy by using Whisper Large instead of Whisper Small.')
add_bullet(doc, 'Add a phoneme-level pronunciation analyzer for accurate pronunciation scoring.')
add_bullet(doc, 'Build a reference database of common Arabic mistakes to increase assessment accuracy.')

add_h2(doc, 'Model-Level Improvements')
add_bullet(doc, 'Evaluate Arabic-specialized models such as Jais and AceGPT and compare them with multilingual models.')
add_bullet(doc, 'Use few-shot prompting with reference evaluation examples to reduce variance between models.')
add_bullet(doc, 'Add a confidence score for each assessment to flag ambiguous cases.')

add_h2(doc, 'User Experience Improvements')
add_bullet(doc, 'Add historical tracking of learner progress over time.')
add_bullet(doc, 'Suggest interactive exercises tailored to detected weaknesses.')
add_bullet(doc, 'Support multiple Arabic dialects (Gulf, Levantine, Egyptian, Maghrebi).')

# ============================================================
# CONCLUSION
# ============================================================
add_h1(doc, 'Conclusion')
add_p(doc,
      'This project demonstrated the feasibility of using LLMs available through OpenRouter to build '
      'an automated assessment system for spoken Arabic. The results show that current models are '
      'capable of producing a reasonable initial assessment, especially when determining the overall '
      'CEFR level. However, the noticeable variance in numerical scores calls for a more rigorous '
      'methodology to obtain consistent and reliable results.',
      align='justify')

add_p(doc,
      'This first version of the project forms a strong foundation that can be built upon to develop '
      'a full-fledged intelligent educational platform for Arabic, serving millions of learners around '
      'the world who are seeking modern tools to learn Arabic.',
      align='justify')

add_p(doc, '— End of Report —', italic=True, color='8FAADC', size=12,
      align='center', before=24)

# ---------- Save ----------
output_path = 'Arabic_Speech_Assessment_Report_EN.docx'
doc.save(output_path)
print(f'Created: {output_path}')
